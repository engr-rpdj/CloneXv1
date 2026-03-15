# backend/ingest.py

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def load_pdf(path: str | None = None, file_bytes: bytes | None = None) -> str:
    """Extract text from a PDF — from file path or raw bytes."""
    from pypdf import PdfReader

    if file_bytes:
        reader = PdfReader(io.BytesIO(file_bytes))
    elif path:
        reader = PdfReader(path)
    else:
        raise ValueError("Provide either path or file_bytes.")

    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + " "
    return text.strip()


def load_text(path: str | None = None, file_bytes: bytes | None = None) -> str:
    """Read plain text — from file path or raw bytes."""
    if file_bytes:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    elif path:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    else:
        raise ValueError("Provide either path or file_bytes.")


def load_docx(path: str | None = None, file_bytes: bytes | None = None) -> str:
    """Extract text from a .docx file — from file path or raw bytes."""
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is required for .docx support. Run: pip install python-docx")

    if file_bytes:
        doc = docx.Document(io.BytesIO(file_bytes))
    elif path:
        doc = docx.Document(path)
    else:
        raise ValueError("Provide either path or file_bytes.")

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}


def load_file(filename: str, file_bytes: bytes) -> str:
    """
    Auto-detect file type from filename and extract text.
    Returns extracted text string.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return load_pdf(file_bytes=file_bytes)
    elif ext in (".txt", ".md"):
        return load_text(file_bytes=file_bytes)
    elif ext == ".docx":
        return load_docx(file_bytes=file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )